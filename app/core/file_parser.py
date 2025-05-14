import PyPDF2
from docx import Document as DocxDocument # Alias to avoid conflict with Langchain Document
from typing import List, Dict, Any
from langchain.docstore.document import Document as LangchainDocument # Langchain's Document
import logging

# to track events that happens in the program while it is executed.
logger = logging.getLogger(__name__)

def parse_pdf(file_path: str, filename: str) -> List[LangchainDocument]:
    docs = []
    try:
        with open(file_path, "rb") as f:
            reader = PyPDF2.PdfReader(f)
            for page_num, page in enumerate(reader.pages):
                text = page.extract_text()
                if text and text.strip():
                    docs.append(LangchainDocument(
                        page_content=text,
                        metadata={"source_document": filename, "page_number": page_num + 1}
                    ))
    except Exception as e:
        logger.error(f"Error parsing PDF {filename}: {e}")
    return docs

def parse_docx(file_path: str, filename: str) -> List[LangchainDocument]:
    docs = []
    try:
        doc = DocxDocument(file_path)
        full_text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        if full_text:
            # For DOCX, page numbers are hard. We can treat the whole doc as one or split by paragraphs.
            # Here, we create one Langchain Document for the whole file.
            # Chunking will happen later.
            docs.append(LangchainDocument(
                page_content=full_text,
                metadata={"source_document": filename, "page_number": None} # Page num not easily available
            ))
    except Exception as e:
        logger.error(f"Error parsing DOCX {filename}: {e}")
    return docs

def parse_txt(file_path: str, filename: str) -> List[LangchainDocument]:
    docs = []
    try:
        with open(file_path, "r", encoding="utf-8") as f:
            text = f.read()
            if text and text.strip():
                docs.append(LangchainDocument(
                    page_content=text,
                    metadata={"source_document": filename, "page_number": None}
                ))
    except Exception as e:
        logger.error(f"Error parsing TXT {filename}: {e}")
    return docs

def parse_file(file_path: str, filename: str) -> List[LangchainDocument]:
    if filename.endswith(".pdf"):
        return parse_pdf(file_path, filename)
    elif filename.endswith(".docx"):
        return parse_docx(file_path, filename)
    elif filename.endswith(".txt"):
        return parse_txt(file_path, filename)
    else:
        logger.warning(f"Unsupported file type: {filename}")
        return []